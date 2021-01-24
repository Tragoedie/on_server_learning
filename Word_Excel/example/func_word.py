from docx import Document
# создаём новый документ:
document = Document()
# добавляем заголовок документа наивысшего уровня
document.add_heading('Document Title', 0)
# добавляем параграф, и запоминаем его в переменной p
p = document.add_paragraph('A plain paragraph having some ')
# текст параграфа можем сразу добавлять
# с различными стилями форматирования:
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
# добавляем заголовок уровнем поменьше
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')
# добавляем параграфы со стилем-списком,
# неупорядоченным и упорядоченным:
document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')
# можно добавить в документ таблицу
# из одной строки и трёх столбцов:
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
document.add_page_break()
# и в заключение сохранить его в файл:
document.save('demo.docx')
